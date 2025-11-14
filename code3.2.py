# Step 1: Install required dependencies (fixed)
!pip install langgraph langchain-openai python-docx pdfplumber openpyxl PyPDF2

import os
import tempfile
from google.colab import files
from langgraph.graph import StateGraph, END
from typing import Dict, Any, List, TypedDict
from langchain_openai import ChatOpenAI

# Document processing imports (fixed)
import pdfplumber
import PyPDF2
from docx import Document
import openpyxl
import io

# Set your OpenAI API key
# Set your OpenAI API key
os.environ["OPENAI_API_KEY"] = "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"


# Step 2: Enhanced Document Processing Pipeline
class DocumentState(TypedDict):
    raw_document: str
    cleaned_text: str
    summary: str
    keywords: List[str]
    processed: bool
    error: str
    file_name: str
    file_type: str
    word_count: int
    page_count: int  # For PDFs

# Step 3: File Upload and Processing Functions (Fixed)
def upload_document():
    """Upload document to Google Colab"""
    print("Please upload a document (PDF, DOCX, TXT, or XLSX):")
    uploaded = files.upload()

    if not uploaded:
        return None, "No file uploaded"

    file_name = list(uploaded.keys())[0]
    file_content = uploaded[file_name]

    return file_name, file_content

def extract_text_from_file(file_name: str, file_content: bytes) -> tuple[str, str, int]:
    """Extract text from various file formats with improved PDF handling"""
    text = ""
    error = None
    page_count = 0

    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_name) as temp_file:
        temp_file.write(file_content)
        temp_path = temp_file.name

    try:
        file_ext = file_name.lower().split('.')[-1]

        if file_ext == 'pdf':
            # Method 1: Using pdfplumber (better for text extraction)
            try:
                with pdfplumber.open(temp_path) as pdf:
                    page_count = len(pdf.pages)
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"--- Page {i+1} ---\n{page_text}\n\n"
            except Exception as e:
                error = f"pdfplumber failed: {str(e)}"

            # Method 2: Fallback to PyPDF2 if pdfplumber fails or no text extracted
            if not text.strip():
                try:
                    with open(temp_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        page_count = len(pdf_reader.pages)
                        for i, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text()
                            if page_text:
                                text += f"--- Page {i+1} ---\n{page_text}\n\n"
                except Exception as e:
                    error = f"PyPDF2 also failed: {str(e)}" if error else f"PyPDF2 failed: {str(e)}"

        elif file_ext == 'docx':
            try:
                doc = Document(temp_path)
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"

                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + " | "
                        text += "\n"
            except Exception as e:
                error = f"Error reading DOCX: {str(e)}"

        elif file_ext in ['xlsx', 'xls']:
            try:
                workbook = openpyxl.load_workbook(temp_path)
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"=== {sheet_name} ===\n"
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join(str(cell) for cell in row if cell is not None)
                        if row_text.strip():
                            text += row_text + "\n"
                    text += "\n"
            except Exception as e:
                error = f"Error reading Excel: {str(e)}"

        elif file_ext == 'txt':
            try:
                # Handle different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        text = file_content.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    error = "Could not decode text file with common encodings"
            except Exception as e:
                error = f"Error reading text file: {str(e)}"

        else:
            error = f"Unsupported file type: .{file_ext}"

    except Exception as e:
        error = f"Unexpected error: {str(e)}"

    finally:
        # Clean up temporary file
        if os.path.exists(temp_path):
            os.unlink(temp_path)

    # If we have an error and no text, return the error
    if error and not text.strip():
        return "", error, 0

    # If we got some text but there were minor errors, still return the text
    return text.strip(), error, page_count

def process_uploaded_document():
    """Main function to handle document upload and processing"""
    file_name, file_content = upload_document()

    if not file_name:
        return None, "No file was uploaded", None, 0

    print(f"📄 Processing file: {file_name}")
    print(f"📊 File size: {len(file_content)} bytes")

    # Extract text from the uploaded file
    extracted_text, error, page_count = extract_text_from_file(file_name, file_content)

    if error and not extracted_text:
        return None, error, None, 0

    if not extracted_text:
        return None, "No text could be extracted from the document", None, 0

    print(f"✅ Successfully extracted {len(extracted_text)} characters")
    if page_count > 0:
        print(f"📑 Pages processed: {page_count}")
    if error:
        print(f"⚠️  Note: {error}")

    file_type = file_name.split('.')[-1].upper() if '.' in file_name else "UNKNOWN"
    return extracted_text, file_name, file_type, page_count

# Step 4: Enhanced Processing Nodes
def validate_input(state: DocumentState):
    """Enhanced validation with file type checking"""
    if not state["raw_document"]:
        return {"error": "No document content provided"}

    if len(state["raw_document"]) < 10:
        return {"error": "Document too short (minimum 10 characters required)"}

    if len(state["raw_document"]) > 50000:  # Reduced for demo purposes
        return {"error": f"Document too large ({len(state['raw_document'])} characters). Maximum 50,000 characters allowed."}

    # Count words for statistics
    word_count = len(state["raw_document"].split())

    return {
        "error": None,
        "word_count": word_count,
        "processed": False
    }

def clean_document(state: DocumentState):
    """Enhanced text cleaning with document structure preservation"""
    text = state["raw_document"]

    # Remove excessive whitespace but preserve some structure
    lines = text.split('\n')
    cleaned_lines = []

    for line in lines:
        line = line.strip()
        if line:  # Only keep non-empty lines
            # Clean up the line but preserve meaningful structure
            cleaned_line = ' '.join(line.split())  # Remove extra spaces within line
            cleaned_lines.append(cleaned_line)

    # Reconstruct with sensible paragraph breaks
    cleaned = '\n'.join(cleaned_lines)

    # Remove common artifacts from document extraction
    cleaning_replacements = [
        ('\x00', ''),  # Remove null characters
        ('\ufeff', ''),  # Remove BOM
        ('\r\n', '\n'),  # Normalize line endings
        ('  ', ' '),  # Replace double spaces with single
    ]

    for old, new in cleaning_replacements:
        cleaned = cleaned.replace(old, new)

    # Truncate if too long for LLM context (keep more context now)
    original_length = len(cleaned)
    if len(cleaned) > 40000:
        cleaned = cleaned[:40000] + f"\n\n[Document truncated from {original_length} to 40000 characters]"

    return {"cleaned_text": cleaned}

def summarize_document(state: DocumentState):
    """Enhanced summarization with context about the document"""
    llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0.3)

    # Prepare context about the document
    file_info = f"File: {state.get('file_name', 'Unknown')} ({state.get('file_type', 'Unknown')})"
    if state.get('page_count', 0) > 0:
        file_info += f", {state['page_count']} pages"
    file_info += f", {state.get('word_count', 0)} words"

    prompt = f"""
    DOCUMENT INFORMATION:
    {file_info}

    DOCUMENT CONTENT:
    {state['cleaned_text']}

    TASK: Create a comprehensive summary of this document.

    Please provide:
    1. MAIN TOPIC: What is the primary subject of this document?
    2. KEY POINTS: What are the 3-5 most important points or findings?
    3. PURPOSE: What is the document's main purpose or conclusion?
    4. KEY DATA: Any important statistics, dates, or facts mentioned?
    5. OVERVIEW: Brief overall summary (2-3 sentences)

    Structure your response clearly with these sections.
    """

    try:
        summary = llm.invoke(prompt)
        return {"summary": summary.content}
    except Exception as e:
        return {"error": f"Failed to generate summary: {str(e)}"}

def extract_keywords(state: DocumentState):
    """Enhanced keyword extraction with categorization"""
    llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0.2)

    prompt = f"""
    DOCUMENT SUMMARY:
    {state['summary']}

    TASK: Extract the most relevant keywords from this document.

    Please provide 5-7 main keywords or key phrases that represent the core topics.
    Return ONLY a comma-separated list of keywords, nothing else.

    Guidelines:
    - Be specific and meaningful
    - Include proper nouns and technical terms
    - Avoid very generic words
    - Focus on the main themes from the summary
    """

    try:
        keywords_response = llm.invoke(prompt)
        keyword_text = keywords_response.content.strip()

        # Clean and split keywords
        keywords = [kw.strip() for kw in keyword_text.split(',')]
        keywords = [kw for kw in keywords if kw and len(kw) > 1]  # Remove empty and single char strings

        # Ensure we have at least some keywords
        if not keywords and state['cleaned_text']:
            # Fallback: simple word-based extraction
            words = state['cleaned_text'].split()
            # Get most frequent meaningful words (longer than 4 chars)
            from collections import Counter
            word_freq = Counter([word.lower() for word in words if len(word) > 4 and word.isalpha()])
            keywords = [word for word, count in word_freq.most_common(5)]

        return {
            "keywords": keywords,
            "processed": True
        }
    except Exception as e:
        return {
            "error": f"Failed to extract keywords: {str(e)}",
            "processed": False
        }

# Step 5: Enhanced Routing and Error Handling
def route_after_validation(state: DocumentState):
    if state.get("error"):
        return "error_handler"
    return "clean"

def error_handler(state: DocumentState):
    error_msg = state.get("error", "Unknown error")
    print(f"❌ Error encountered: {error_msg}")
    return {"processed": False}

def success_handler(state: DocumentState):
    """Node to handle successful processing completion"""
    print("✅ Document processing completed successfully!")
    return {"processed": True}

# Step 6: Build the Enhanced Graph
builder = StateGraph(DocumentState)

# Add nodes
builder.add_node("validate", validate_input)
builder.add_node("clean", clean_document)
builder.add_node("summarize", summarize_document)
builder.add_node("extract_keywords", extract_keywords)
builder.add_node("error_handler", error_handler)
builder.add_node("success_handler", success_handler)

# Define workflow
builder.set_entry_point("validate")
builder.add_conditional_edges("validate", route_after_validation)
builder.add_edge("clean", "summarize")
builder.add_edge("summarize", "extract_keywords")
builder.add_edge("extract_keywords", "success_handler")
builder.add_edge("success_handler", END)
builder.add_edge("error_handler", END)

# Compile the graph
graph = builder.compile()

# Step 7: Main Execution Function
def process_document_pipeline():
    """Complete document processing pipeline"""
    print("📁 Document Processing Pipeline")
    print("=" * 50)

    # Upload and extract document
    result = process_uploaded_document()
    extracted_text, file_name, file_type, page_count = result

    if not extracted_text:
        print(f"❌ Failed: {file_name}")  # file_name contains error message in this case
        return None

    # Prepare initial state
    initial_state = {
        "raw_document": extracted_text,
        "cleaned_text": "",
        "summary": "",
        "keywords": [],
        "processed": False,
        "error": "",
        "file_name": file_name,
        "file_type": file_type,
        "word_count": 0,
        "page_count": page_count
    }

    print(f"🔄 Processing document through LangGraph pipeline...")

    # Process through the graph
    try:
        result = graph.invoke(initial_state)

        # Display results
        print("\n" + "=" * 50)
        print("📊 PROCESSING RESULTS")
        print("=" * 50)
        print(f"📄 File: {result.get('file_name', 'N/A')}")
        print(f"📝 Type: {result.get('file_type', 'N/A')}")
        if result.get('page_count'):
            print(f"📑 Pages: {result.get('page_count')}")
        print(f"📊 Word Count: {result.get('word_count', 0)}")
        print(f"✅ Processing Status: {'Success' if result.get('processed') else 'Failed'}")

        if result.get("processed"):
            print(f"\n📋 SUMMARY:")
            print("-" * 30)
            print(result.get("summary", "No summary generated"))

            print(f"\n🔑 KEYWORDS:")
            print("-" * 30)
            keywords = result.get("keywords", [])
            for i, keyword in enumerate(keywords, 1):
                print(f"  {i}. {keyword}")
        else:
            print(f"\n❌ Error: {result.get('error', 'Unknown error')}")

        return result

    except Exception as e:
        print(f"❌ Pipeline execution failed: {str(e)}")
        return None

# Step 8: Alternative - Process Sample Document
def process_sample_document():
    """Process a sample document without upload"""
    sample_doc = """
    ARTIFICIAL INTELLIGENCE IN HEALTHCARE: CURRENT APPLICATIONS AND FUTURE PROSPECTS

    Executive Summary:
    This report examines the transformative impact of artificial intelligence on the healthcare industry.
    AI technologies are revolutionizing medical diagnostics, treatment planning, and operational efficiency.

    Key Findings:

    1. MEDICAL IMAGING AND DIAGNOSTICS
    - AI systems can analyze X-rays with 95% accuracy compared to 92% for human radiologists
    - Machine learning models detect early-stage cancers with 87% sensitivity
    - Reduction in diagnostic errors: up to 85% improvement in certain specialties

    2. DRUG DISCOVERY AND DEVELOPMENT
    - AI accelerated drug discovery process by 40% in recent trials
    - Predictive models identify promising drug candidates with 75% accuracy
    - Clinical trial optimization reduces development timelines by 30%

    3. PERSONALIZED MEDICINE
    - Genetic analysis algorithms create customized treatment plans
    - Predictive analytics identify patient-specific risk factors
    - Treatment success rates improved by 25% through personalization

    4. OPERATIONAL EFFICIENCY
    - Administrative automation reduces paperwork by 60%
    - Patient scheduling optimization decreases wait times by 45%
    - Resource allocation algorithms cut operational costs by 35%

    Challenges and Considerations:
    - Data privacy and security concerns remain significant
    - Algorithm transparency and explainability need improvement
    - Integration with legacy healthcare systems presents technical hurdles
    - Regulatory compliance requires careful attention

    Future Outlook:
    The global AI in healthcare market is projected to reach $45.2 billion by 2026,
    growing at a CAGR of 44.9%. Emerging applications include robotic surgery assistance,
    virtual health assistants, and predictive outbreak monitoring.

    Conclusion:
    AI represents a paradigm shift in healthcare delivery, offering unprecedented opportunities
    for improved patient outcomes and operational excellence. Successful implementation requires
    careful consideration of ethical, technical, and regulatory factors.
    """

    initial_state = {
        "raw_document": sample_doc,
        "cleaned_text": "",
        "summary": "",
        "keywords": [],
        "processed": False,
        "error": "",
        "file_name": "sample_ai_healthcare_report.pdf",
        "file_type": "PDF",
        "word_count": 0,
        "page_count": 3
    }

    print("🧪 Processing Sample Document...")
    result = graph.invoke(initial_state)

    print("\n📊 SAMPLE DOCUMENT RESULTS")
    print("=" * 40)
    print(f"✅ Processing complete: {result['processed']}")
    print(f"\n📋 Summary:")
    print("-" * 20)
    print(result['summary'])
    print(f"\n🔑 Keywords: {', '.join(result['keywords'])}")

    return result

# Step 9: Usage Examples
print("🚀 Enhanced Document Processing Pipeline Ready!")
print("\nChoose an option:")
print("1. process_document_pipeline() - Upload and process your own document")
print("2. process_sample_document() - Process the built-in sample document")

# Uncomment the line below to test with sample document
# process_sample_document()

# Uncomment the line below to upload and process your own document
# process_document_pipeline()
